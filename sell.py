from ast import main
import tkinter as tk
from datetime import datetime
from tkinter import ttk
import smtplib
from docx import Document
import docx
from email.message import EmailMessage
from styles import *
from tkinter import PhotoImage
from PIL import Image, ImageTk

import re
from sqlite3 import Cursor
from tkinter import simpledialog
import shared_variables
import pandas as pd
from tkinter import filedialog
import mysql.connector
import random
import string
from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from time import strftime
from datetime import date
def sell(root):
    def display_data(tree):
        connection = mysql.connector.connect(
            host='localhost',
            user='root',
            password='080102',
            database='myDatabase'
        )
        cursor = connection.cursor()
        cursor.execute("SELECT product_name, category, stock, product_price FROM inventory WHERE stock>0")
        fetch = cursor.fetchall()
        
        # Xóa tất cả các cột trước khi thêm mới để tránh việc trùng lặp dữ liệu
        tree.delete(*tree.get_children())
        for data in fetch:
            tree.insert("", "end", values=data)
        tree.all_items = fetch

    def search_product():
        connection = mysql.connector.connect(
            host='localhost',
            user='root',
            password='080102',
            database='myDatabase'
        )
        cursor = connection.cursor()
        cursor.execute("SELECT product_name, category, stock, product_price FROM inventory WHERE stock>0")
        fetch = cursor.fetchall()
        for data in fetch:
            tree.insert("", "end", values=data)
        all_items = fetch
        search_text = search_entry.get()  # Lấy giá trị từ ô tìm kiếm
        if not search_text:
            for item in tree.get_children():
                tree.delete(item)
        # Hiển thị toàn bộ sản phẩm trên treeview
            for item in all_items:
                tree.insert("", "end", values=item)
            return
        # Xóa toàn bộ sản phẩm hiển thị trên treeview
        for item in tree.get_children():
            tree.delete(item)

        found = False
        for item in all_items:  # all_items là danh sách tất cả các sản phẩm trước khi tìm kiếm
            if search_text in str(item[0]):  # Kiểm tra xem giá trị tìm kiếm có trong sản phẩm hay không
                tree.insert("", "end", values=item)  # Nếu tìm thấy, thêm sản phẩm đó vào treeview
                found = True

            if search_text in str(item[1]):  # Kiểm tra xem giá trị tìm kiếm có trong sản phẩm hay không
                tree.insert("", "end", values=item)  # Nếu tìm thấy, thêm sản phẩm đó vào treeview
                found = True

            if search_text in str(item[2]):  # Kiểm tra xem giá trị tìm kiếm có trong sản phẩm hay không
                tree.insert("", "end", values=item)  # Nếu tìm thấy, thêm sản phẩm đó vào treeview
                found = True

        if not found:
            messagebox.showerror("Oops!!", f"Product ID: {search_text} not found.")

    #gọi hàm style căn chỉnh
    configure_styles()
    sell_window = tk.Toplevel(root)
    sell_window.title("sell")
        # Lấy kích thước của màn hình
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    #1366x768
    window_width = 1366  # Thay đổi kích thước theo nhu cầu
    window_height = 768  # Thay đổi kích thước theo nhu cầu
    # Tính toán vị trí để cửa sổ xuất hiện giữa màn hình
    x = (screen_width - window_width) // 2
    y = (screen_height - window_height) // 2

    # Đặt vị trí cửa sổ
    sell_window.geometry(f"{window_width}x{window_height}+{x}+{y}")
    sell_window.resizable(width=False, height=False)

    def close_window_2():
        sell_window.destroy()  # Đóng cửa sổ 2
        root.deiconify() 
    # frame_back = ttk.Frame(sell_window)
    # frame_back.place(x=20,y=20)
    back_button = ttk.Button(sell_window, text="Quay lại", command=close_window_2, style='Back_Bbutton.TButton')
    back_button.place(relx=0.007, rely=0.006, width=100, height=45)
    font_style = ('Arial', 24,'bold')
    label_shop = ttk.Label(sell_window, text="Happy Shop", foreground="green", font=font_style)
    label_shop.place(relx=0.4, rely=0.004)
    
    scrollbarx = Scrollbar(sell_window, orient=HORIZONTAL)
    scrollbary = Scrollbar(sell_window, orient=VERTICAL)
    tree = ttk.Treeview(sell_window)
    tree.place(relx=0.03, rely=0.2, width=850, height=550)
    tree.configure(
            yscrollcommand=scrollbary.set, xscrollcommand=scrollbarx.set
        )
    tree.configure(selectmode="extended")

    # tree.bind("<<TreeviewSelect>>", on_tree_select)

    scrollbary.configure(command=tree.yview)
    scrollbarx.configure(command=tree.xview)

    scrollbary.place(relx=0.65, rely=0.2, width=22, height=548)
    scrollbarx.place(relx=0.03, rely=0.92, width=884, height=22)
    tree.configure(
            columns=(
                "Tên sản phẩm",
                "Loại sản phẩm",
                "Số lượng",
                "Giá",
            )
        )
#product_id, product_name,category ,stock, product_price

    tree.heading("Tên sản phẩm", text="Tên sản phẩm", anchor=W)
    tree.heading("Loại sản phẩm", text="Loại sản phẩm", anchor=W)
    tree.heading("Số lượng", text="Số lượng", anchor=W)
    tree.heading("Giá", text="Giá tiền (VND)", anchor=W)



    tree.column("#0", stretch=NO, minwidth=0, width=2)
    tree.column("#1", stretch=NO, minwidth=0, width=280)
    tree.column("#2", stretch=NO, minwidth=0, width=260)
    tree.column("#3", stretch=NO, minwidth=0, width=165)
    tree.column("#4", stretch=NO, minwidth=0, width=165)

    # Tạo treeview_selected
    treeview_selected = ttk.Treeview(sell_window)
    treeview_selected.place(relx=0.68, rely=0.2, width=380, height=400)
    # Cấu hình các cột cho treeview_selected
    treeview_selected.configure(
        columns=(
            "Tên sản phẩm",
            "Loại sản phẩm",
            "Số lượng",
            "Đơn Giá"
        )
    )

    # Đặt tiêu đề cho các cột
    treeview_selected.heading("Tên sản phẩm", text="Tên sản phẩm", anchor=tk.W)
    treeview_selected.heading("Loại sản phẩm", text="Loại sản phẩm", anchor=tk.W)
    treeview_selected.heading("Số lượng", text="Số lượng", anchor=tk.W)
    treeview_selected.heading("Đơn Giá", text="Đơn Giá", anchor=tk.W)

    # Cấu hình các cột cho treeview_selected
    treeview_selected.column("#0", stretch=tk.NO, minwidth=0, width=0)
    treeview_selected.column("#1", stretch=tk.NO, minwidth=0, width=100)
    treeview_selected.column("#2", stretch=tk.NO, minwidth=0, width=100)
    treeview_selected.column("#3", stretch=tk.NO, minwidth=0, width=100)
    treeview_selected.column("#4", stretch=tk.NO, minwidth=0, width=79)

    def on_tree_select(event):
        selected_items = tree.selection()
        for item in selected_items:
            values = tree.item(item, "values")
            product_name = values[0]
            product_type = values[1]
            price = values[3]

            # Kiểm tra xem sản phẩm đã tồn tại trong treeview_selected chưa
            item_exists = False
            for child in treeview_selected.get_children():
                values_selected = treeview_selected.item(child, "values")
                selected_name = values_selected[0]
                selected_type = values_selected[1]
                if product_name == selected_name and product_type == selected_type:
                    # Sản phẩm đã tồn tại, cộng số lượng
                    current_stock = int(values_selected[2])
                    new_stock = current_stock + 1
                    treeview_selected.set(child, "Số lượng", new_stock)
                    item_exists = True
                    break

            if not item_exists:
                # Thêm sản phẩm mới vào treeview_selected
                treeview_selected.insert("", "end", values=(product_name, product_type, 1, price))
    def is_valid_input(char):
        return char.isdigit() or char == ""
    # Tạo Label với kiểu đã đặt
    name_label = ttk.Label(sell_window, text="Tên khách hàng", style="Round.TLabel")
    name_label.place(relx=0.38, rely=0.1, width=115, height=30)
    name_entry = ttk.Entry(sell_window)
    name_entry.place(relx=0.47, rely=0.1,  width=130, height=30)

    phone_label = ttk.Label(sell_window, text="Số điện thoại", style="Round.TLabel")
    phone_label.place(relx=0.61, rely=0.1, width=100, height=30)
    # phone_label = ttk.Label(sell_window, text="Phone:")
    # phone_label.place(relx=0.61, rely=0.1, width=100, height=30)

    validate_cmd = (sell_window.register(lambda char: is_valid_input(char)), '%S')
    phone_entry = ttk.Entry(sell_window, validate="key", validatecommand=validate_cmd)
    phone_entry.place(relx=0.69, rely=0.1, width=120, height=30)
    

    email_label = ttk.Label(sell_window, text="Email", style="Round.TLabel")
    email_label.place(relx=0.82, rely=0.1, width=60, height=30)
    email_entry = ttk.Entry(sell_window)
    email_entry.place(relx=0.87, rely=0.1,  width=120, height=30)
    # Gắn sự kiện "<<TreeviewSelect>>" cho tree
    tree.bind("<<TreeviewSelect>>", on_tree_select)
    # Tạo hàm xử lý sự kiện khi nhấn nút "Thanh toán"
    connection = mysql.connector.connect(
            host='localhost',
            user='root',
            password='080102',
            database='myDatabase'
    )
    cursor = connection.cursor()
    def calculate_total():
        customer_name = name_entry.get()
        phone_number = phone_entry.get()
        email = email_entry.get()
        purchase_datetime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        total = 0
        selected_items = treeview_selected.get_children()
        invoice_items = []  # Danh sách để lưu chi tiết các món hàng trong hoá đơn

        khuyenmai = float(khuyenmai_combobox.get().strip("%")) / 100
        for item in selected_items:
            product_name = treeview_selected.item(item, "values")[0]  # Tên sản phẩm là cột đầu tiên
            quantity = float(treeview_selected.item(item, "values")[2])  # Số lượng là cột thứ 3
            price = float(treeview_selected.item(item, "values")[3])  # Đơn giá là cột thứ 4
            subtotal = quantity * price
            subtotal_with_discount = subtotal -subtotal*khuyenmai
            total += subtotal_with_discount

            # Thêm chi tiết món hàng vào danh sách
            invoice_items.append((product_name, quantity, price, subtotal))

        # Lưu thông tin hoá đơn vào bảng invoices
        sql_invoice = "INSERT INTO invoices (total, customer_name, phone_number, email) VALUES (%s, %s, %s, %s)"
        values_invoice = (total, customer_name, phone_number, email)
        cursor.execute(sql_invoice, values_invoice)
        connection.commit()


        # Lấy ID của hoá đơn vừa được thêm vào
        invoice_id = cursor.lastrowid

        # Lưu chi tiết hoá đơn vào bảng invoice_details
        sql_invoice_detail = "INSERT INTO invoice_details (invoice_id, product_name, quantity, price, subtotal, customer_name, purchase_datetime) VALUES (%s, %s, %s, %s, %s, %s, %s)"
        values_invoice_detail = [(invoice_id, item[0], item[1], item[2], item[3], customer_name, purchase_datetime) for item in invoice_items]
        cursor.executemany(sql_invoice_detail, values_invoice_detail)
        connection.commit()

        # Cập nhật số lượng trong cơ sở dữ liệu
        for item in selected_items:
            product_name = treeview_selected.item(item, "values")[0]
            quantity = float(treeview_selected.item(item, "values")[2])

            # Lấy số lượng ban đầu từ cơ sở dữ liệu
            sql_get_initial_quantity = "SELECT stock FROM inventory WHERE product_name = %s"
            cursor.execute(sql_get_initial_quantity, (product_name,))
            initial_quantity = cursor.fetchone()[0]

            # Tính toán số lượng mới
            new_quantity = initial_quantity - quantity

            # Cập nhật số lượng mới trong cơ sở dữ liệu
            sql_update_quantity = "UPDATE inventory SET stock = %s WHERE product_name = %s"
            cursor.execute(sql_update_quantity, (new_quantity, product_name))
            connection.commit()


        # Hiển thị lại dữ liệu trên display_data
        display_data(tree)
        
        # Hiển thị giá tổng đơn hàng bằng hộp thoại thông báo
        messagebox.showinfo("Thông báo", f"Tổng đơn hàng: {total} đồng")
        treeview_selected.insert("", "end", text="")
        treeview_selected.insert("", "end",  values=["Thành tiền", "", "", f"{total} đồng"])

        



    #tao nut ship hang


    #tạo nút "hàng hoá giảm giá"

   



    # button_deliver = ttk.Button(sell_window, text="Giao hàng tại nhà")
    # button_deliver.place(relx=0.58, rely=0.16)

    # Tạo hàm xử lý sự kiện khi nhấn nút "Xóa"
    def delete_all_data():
        treeview_selected.delete(*treeview_selected.get_children())

    # Tạo nút "Xóa"
    button_delete = ttk.Button(sell_window, text="Xóa", command=delete_all_data)
    button_delete.place(relx=0.68, rely=0.16)


    display_data(tree)
   
    search_entry = ttk.Entry(sell_window)
    search_entry.place(relx=0.03, rely=0.1, width=200, height=30)
    button1 = Button(sell_window)
    button1.place(relx=0.19, rely=0.105, width=76, height=28)
    button1.configure(relief="flat")
    button1.configure(overrelief="flat")
    button1.configure(activebackground="#CF1E14")
    button1.configure(cursor="hand2")
    button1.configure(foreground="#ffffff")
    button1.configure(background="#CF1E14")
    button1.configure(font="-family {Poppins SemiBold} -size 10")
    button1.configure(borderwidth="0")
    button1.configure(text="""Tìm kiếm""")
    button1.configure(command=search_product)
    style = ttk.Style()
    style.configure("Round.TLabel", foreground="white", background="green", borderwidth=2, relief="solid", padding=(10, 5))
    def sendEmail():
        EMAIL = 'trieutungvp@outlook.com'
        PASSWORD = 'trieuthanhtung@0801'
        DESTINATION_EMAIL = email_entry.get()
        SUBJECT_EMAIL = "Tiêu đề: Lập trình mạng"
        BODY_EMAIL = "Cảm ơn bạn đã mua hàng tại đây"
        msg = EmailMessage()
        msg['To'] = DESTINATION_EMAIL
        msg['From'] = EMAIL
        msg['Subject'] = SUBJECT_EMAIL
        msg.set_content(BODY_EMAIL)
        mailServer = 'smtp.office365.com'
        mailPort = 587
        connection = smtplib.SMTP(mailServer, mailPort)
        connection.starttls()
        connection.login(EMAIL, PASSWORD)
        connection.send_message(msg=msg, from_addr=EMAIL, to_addrs=DESTINATION_EMAIL)
        connection.quit()

    
    def edit_quantity():
        # Lấy item được chọn trong treeview_selected
        selected_item = treeview_selected.focus()

        if selected_item:
            # Lấy giá trị của cột "Số lượng" của item đang được chọn
            current_quantity = treeview_selected.item(selected_item, "values")[2]

            # Tạo một cửa sổ popup cho phép chỉnh sửa số lượng
            popup_window = tk.Toplevel(sell_window)

            # Tính toán vị trí giữa màn hình cho cửa sổ popup
            window_width = 200
            window_height = 150
            screen_width = sell_window.winfo_screenwidth()
            screen_height = sell_window.winfo_screenheight()
            x = (screen_width - window_width) // 2
            y = (screen_height - window_height) // 2

            # Đặt tọa độ cửa sổ popup
            popup_window.geometry(f"{window_width}x{window_height}+{x}+{y}")

            # Tạo một nhãn và ô văn bản để hiển thị và chỉnh sửa số lượng
            quantity_label = ttk.Label(popup_window, text="Số lượng:")
            quantity_label.pack()

            quantity_entry = ttk.Entry(popup_window)
            quantity_entry.insert(0, current_quantity)
            quantity_entry.pack()

            # Tạo một hàm xử lý sự kiện khi nhấn nút "Lưu"
            def save_quantity():
                new_quantity = quantity_entry.get()
                
                if not new_quantity.isdigit() or int(new_quantity) <= 0:
                    messagebox.showerror("Lỗi", "Số lượng phải là một số nguyên dương.")
                    return
                new_quantity = int(new_quantity)
                # Lấy item được chọn trong treeview_selected
                selected_item = treeview_selected.focus()

                # Kết nối đến cơ sở dữ liệu MySQL
                connection = mysql.connector.connect(
                    host='localhost',
                    user='root',
                    password='080102',
                    database='myDatabase'
                )
                cursor = connection.cursor()
                product_name = treeview_selected.item(selected_item, "values")[0]

                # Truy vấn số lượng sản phẩm từ bảng inventory
                query = "SELECT stock FROM inventory WHERE product_name= %s"
                cursor.execute(query, (product_name,))
                result = cursor.fetchone()
                
                new_quantity = int(new_quantity)
                if new_quantity>result[0]:
                    messagebox.showwarning("Lỗi", "Số lượng không đủ.")
                    return
                # Cập nhật giá trị số lượng trong treeview_selected
                treeview_selected.set(selected_item, "Số lượng", new_quantity)

                # Đóng cửa sổ popup
                popup_window.destroy()

            # Tạo nút "Lưu" để lưu giá trị mới của số lượng
            save_button = ttk.Button(popup_window, text="Lưu", command=save_quantity)
            save_button.pack()
    edit_button = ttk.Button(sell_window, text="Chỉnh số lượng", command=edit_quantity)
    edit_button.place(relx=0.82, rely=0.16)

    send_email_button = ttk.Button(sell_window, text='Gửi email',command=sendEmail)
    send_email_button.place(relx=0.9,rely=0.16)




    label_payment = ttk.Label(sell_window, text="Khuyến mãi", font=("Arial", 12, "bold"))
    label_payment.place(relx=0.68, rely=0.75)
    # Tạo combobox
    khuyenmai_combobox = ttk.Combobox(sell_window, values=["0%", "5%", "10%", "15%"],state="readonly")
    # Đặt giá trị mặc định cho combobox
    khuyenmai_combobox.set("0%")
    # Định vị trí và kích thước của combobox
    khuyenmai_combobox.place(relx=0.68, rely=0.8, width=100, height=25)
    # Tạo nút "Thanh toán"
    button_payment = ttk.Button(sell_window, text="Thanh toán", command=calculate_total)
    button_payment.place(relx=0.68, rely=0.86, width=100, height=25)
    def display_zero_stock():
        connection = mysql.connector.connect(
            host='localhost',
            user='root',
            password='080102',
            database='myDatabase'
        )
        cursor = connection.cursor()
        cursor.execute("SELECT product_name, category, stock, product_price FROM inventory WHERE stock = 0")
        fetch = cursor.fetchall()
        
        window = tk.Toplevel(sell_window)  # Tạo cửa sổ con
        window.title("Sản phẩm hết hàng")
        treeview = ttk.Treeview(window, columns=("Product Name", "Category", "Stock", "Price"), show="headings")
        treeview.heading("Product Name", text="Tên Sản Phẩm")
        treeview.heading("Category", text="Loại Sản Phẩm")
        treeview.heading("Stock", text="Số Lượng")
        treeview.heading("Price", text="Giá")
        
        for data in fetch:
            treeview.insert("", "end", values=data)
        
        treeview.pack()

    # Tạo nút "Sản phẩm hết hàng"
    button_display_zero = ttk.Button(sell_window, text="Sản phẩm hết hàng", command=display_zero_stock)
    button_display_zero.place(relx=0.03, rely=0.16)

    def sort_by_product_name():
       
        tree_data = [(tree.set(item, "Tên sản phẩm"), item) for item in tree.get_children()]
        sorted_data = sorted(tree_data, key=lambda x: x[0])
        for index, (product_name, item) in enumerate(sorted_data):
            tree.move(item, '', index)

 
    tree.heading("Tên sản phẩm", text="Tên sản phẩm", anchor=tk.W, command=sort_by_product_name)

    def sort_by_category():
       
        tree_data = [(tree.set(item, "Loại sản phẩm"), item) for item in tree.get_children()]
        sorted_data = sorted(tree_data, key=lambda x: x[0])
        for index, (category, item) in enumerate(sorted_data):
            tree.move(item, '', index)

  
    tree.heading("Loại sản phẩm", text="Loại sản phẩm", anchor=tk.W, command=sort_by_category)

    def sort_by_stock():
  
        tree_data = [(int(tree.set(item, "Số lượng")), item) for item in tree.get_children()]
        sorted_data = sorted(tree_data, key=lambda x: x[0])
        for index, (stock, item) in enumerate(sorted_data):
            tree.move(item, '', index)

    tree.heading("Số lượng", text="Số lượng", anchor=tk.W, command=sort_by_stock)

    def sort_by_product_price():
       
        tree_data = [(int(tree.set(item, "Giá")), item) for item in tree.get_children()]
        sorted_data = sorted(tree_data, key=lambda x: x[0])
        for index, (product_price, item) in enumerate(sorted_data):
            tree.move(item, '', index)

    tree.heading("Giá", text="Giá tiền (VND)", anchor=tk.W, command=sort_by_product_price)

    def open_word_file():
        # Đường dẫn tới file Word "loinhan.docx"
        file_path = "loinhan.docx"
        
        # Đọc dữ liệu từ file Word
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Hiển thị dữ liệu trong cửa sổ mới
        new_window = tk.Toplevel()
        new_window.title("Lịch sử lời nhắn")
        text_widget = tk.Text(new_window)
        text_widget.insert(tk.END, text)
        text_widget.configure(state='disabled')  # Chỉ cho phép đọc
        text_widget.pack()
        # Hiển thị emp_id trong Text widget
        text_widget.insert(tk.END, f"Emp ID: {shared_variables.emp_id}\n")
        # Tạo nút "Ghi lại"
        button_save = ttk.Button(new_window, text="Ghi lại", command=lambda: open_save_window(new_window))
        button_save.pack()

    def open_save_window(parent_window):
        # Tạo cửa sổ con mới
        save_window = tk.Toplevel(parent_window)
        save_window.title("Ghi lời nhắn")
        # Tạo Text widget cho việc nhập liệu
        text_widget = tk.Text(save_window)
        text_widget.pack()
        
        # Tạo nút "Lưu"
        button_save = ttk.Button(save_window, text="Lưu", command=lambda: save_message(text_widget, save_window))
        button_save.pack()

    def save_message(text_widget, save_window):
        # Đường dẫn tới file Word "loinhan.docx"
        file_path = "loinhan.docx"
        
        # Tạo đối tượng Document từ file Word
        doc = Document(file_path)
        
        # Lấy ngày và giờ hiện tại
        now = datetime.now()
        timestamp = now.strftime("%d/%m/%Y %H:%M:%S")
        
        # Lấy nội dung tin nhắn từ Text widget
        message = text_widget.get("1.0", tk.END)
        
        # Ghi lại lời nhắn vào tài liệu
        doc.add_paragraph(f"<{timestamp}> <{shared_variables.emp_id}>: {message}")
        
        # Lưu tài liệu
        doc.save(file_path)
        
        messagebox.showinfo("Thông báo", f"Lưu tin nhắn thành công")

        # Đóng cửa sổ con
        save_window.destroy()

    # Tạo nút "Lời nhắn"
    button_notification = ttk.Button(sell_window, text="Lời nhắn", command=open_word_file)
    button_notification.place(relx=0.14, rely=0.16)

